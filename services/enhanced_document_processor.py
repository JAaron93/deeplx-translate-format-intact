"""
Enhanced document processor with comprehensive formatting preservation.
Integrates advanced PDF processing with other document formats.
"""

from __future__ import annotations

import os
import json
import logging
import tempfile
import asyncio
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
import fitz
from docx import Document
import mammoth

from .advanced_pdf_processor import AdvancedPDFProcessor, PageLayout
from .dolphin_client import get_layout as get_dolphin_layout

logger = logging.getLogger(__name__)


def validate_dolphin_layout(layout: Dict[str, Any], expected_page_count: int) -> bool:
    """Validate the structure of the Dolphin layout data.
    
    Args:
        layout: The Dolphin layout data to validate
        expected_page_count: Expected number of pages in the layout
        
    Returns:
        bool: True if layout is valid, False otherwise
    """
    if not isinstance(layout, dict):
        logger.warning("Dolphin layout must be a dictionary")
        return False
        
    if 'pages' not in layout:
        logger.warning("Dolphin layout missing 'pages' key")
        return False
        
    if not isinstance(layout['pages'], list):
        logger.warning("Dolphin layout 'pages' must be a list")
        return False
        
    if len(layout['pages']) != expected_page_count:
        logger.warning(
            "Dolphin layout page count mismatch. "
            f"Expected {expected_page_count}, got {len(layout['pages'])}"
        )
        return False
        
    # Validate each page structure
    for i, page in enumerate(layout['pages']):
        if not isinstance(page, dict):
            logger.warning(f"Page {i} is not a dictionary")
            return False
            
        # Add more specific validations here based on Dolphin's schema
        # For example, check for required fields in each page
        
    return True

@dataclass
class DocumentMetadata:
    """Metadata for processed documents."""
    filename: str
    file_type: str
    total_pages: int
    total_text_elements: int
    file_size_mb: float
    processing_time: float
    dpi: int = 300

class EnhancedDocumentProcessor:
    """
    Enhanced document processor with comprehensive formatting preservation.
    Supports PDF, DOCX, and TXT with advanced layout preservation for PDFs.
    """
    
    def __init__(self, dpi: int = 300, preserve_images: bool = True) -> None:
        """
        Initialize the enhanced document processor.
        
        Args:
            dpi: Resolution for PDF processing
            preserve_images: Whether to preserve images in PDFs
        """
        self.dpi = dpi
        self.preserve_images = preserve_images
        self.pdf_processor: AdvancedPDFProcessor = AdvancedPDFProcessor(dpi=dpi, preserve_images=preserve_images)
        
    def extract_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extract content from document with format-specific processing.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        logger.info(f"Processing document: {file_path} ({file_ext})")
        
        if file_ext == '.pdf':
            return self._extract_pdf_content(file_path)
        elif file_ext == '.docx':
            return self._extract_docx_content(file_path)
        elif file_ext == '.txt':
            return self._extract_txt_content(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_pdf_content(self, pdf_path: str) -> Dict[str, Any]:
        """Extract content from PDF with advanced layout preservation."""
        import time
        start_time = time.time()
        
        # Extract complete layout information using internal processor
        layouts = self.pdf_processor.extract_document_layout(pdf_path)

        # Retrieve Dolphin layout for higher-fidelity page structure (optional)
        try:
            try:
                # Try to get the current event loop
                loop = asyncio.get_running_loop()
                # If we're in an async context, we need to run in a thread
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(asyncio.run, get_dolphin_layout(pdf_path))
                    dolphin_layout = future.result()
            except RuntimeError:
                # No running loop, safe to use asyncio.run
                dolphin_layout = asyncio.run(get_dolphin_layout(pdf_path))
        except Exception as dl_err:
            logger.warning("Failed to fetch Dolphin layout for %s: %s", pdf_path, dl_err)
            dolphin_layout = None
        # Create backup of layout information
        backup_path = self._get_backup_path(pdf_path)
        self.pdf_processor.create_layout_backup(layouts, backup_path)
        
        # Extract text for translation
        text_by_page = self.pdf_processor.get_text_for_translation(layouts)
        
        # Calculate metadata
        total_text_elements = sum(len(layout.text_elements) for layout in layouts)
        file_size_mb = os.path.getsize(pdf_path) / (1024 * 1024)
        processing_time = time.time() - start_time
        
        metadata = DocumentMetadata(
            filename=Path(pdf_path).name,
            file_type='PDF',
            total_pages=len(layouts),
            total_text_elements=total_text_elements,
            file_size_mb=file_size_mb,
            processing_time=processing_time,
            dpi=self.dpi
        )
        
        # Validate dolphin_layout structure if present
        if dolphin_layout and not validate_dolphin_layout(dolphin_layout, len(layouts)):
            logger.warning("Invalid Dolphin layout structure, discarding")
            dolphin_layout = None
        return {
            'type': 'pdf_advanced',
            'layouts': layouts,
            'text_by_page': text_by_page,
            'metadata': metadata,
            'backup_path': backup_path,
            'preview': self.pdf_processor.generate_preview(layouts),
            'dolphin_layout': dolphin_layout
        }
    
    def _extract_docx_content(self, docx_path: str) -> Dict[str, Any]:
        """Extract content from DOCX file."""
        import time
        start_time = time.time()
        
        try:
            # Extract text using mammoth for better formatting preservation
            with open(docx_path, 'rb') as docx_file:
                result = mammoth.extract_raw_text(docx_file)
                text_content = result.value
                
            # Also extract using python-docx for structure information
            doc = Document(docx_path)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            file_size_mb = os.path.getsize(docx_path) / (1024 * 1024)
            processing_time = time.time() - start_time
            
            metadata = DocumentMetadata(
                filename=Path(docx_path).name,
                file_type='DOCX',
                total_pages=1,  # DOCX doesn't have fixed pages
                total_text_elements=len(paragraphs),
                file_size_mb=file_size_mb,
                processing_time=processing_time
            )
            
            return {
                'type': 'docx',
                'text_content': text_content,
                'paragraphs': paragraphs,
                'metadata': metadata,
                'preview': text_content[:1000] + '...' if len(text_content) > 1000 else text_content
            }
            
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            logger.error(
                f"Failed to process DOCX file '{Path(docx_path).name}'. "
                f"Error during {'text extraction' if 'mammoth' in error_details else 'paragraph processing' if 'paragraphs' in error_details else 'metadata generation'}: {str(e)}\n"
                f"Traceback:\n{error_details}"
            )
            raise
    
    def _extract_txt_content(self, txt_path: str) -> Dict[str, Any]:
        """Extract content from TXT file."""
        import time
        start_time = time.time()
        
        try:
            with open(txt_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Split into lines for processing
            lines = [line for line in text_content.split('\n') if line.strip()]
            
            file_size_mb = os.path.getsize(txt_path) / (1024 * 1024)
            processing_time = time.time() - start_time
            
            metadata = DocumentMetadata(
                filename=Path(txt_path).name,
                file_type='TXT',
                total_pages=1,
                total_text_elements=len(lines),
                file_size_mb=file_size_mb,
                processing_time=processing_time
            )
            
            return {
                'type': 'txt',
                'text_content': text_content,
                'lines': lines,
                'metadata': metadata,
                'preview': text_content[:1000] + '...' if len(text_content) > 1000 else text_content
            }
            
        except Exception as e:
            logger.error(f"Error processing TXT file: {e}")
            raise
    
    def create_translated_document(self, original_content: Dict[str, Any], 
                                 translated_texts: Dict[int, List[str]], 
                                 output_filename: str) -> str:
        """Create translated document preserving original formatting."""
        
        output_path = os.path.join('downloads', output_filename)
        os.makedirs('downloads', exist_ok=True)
        
        content_type = original_content['type']
        
        if content_type == 'pdf_advanced':
            return self._create_translated_pdf(original_content, translated_texts, output_path)
        elif content_type == 'docx':
            return self._create_translated_docx(original_content, translated_texts, output_path)
        elif content_type == 'txt':
            return self._create_translated_txt(original_content, translated_texts, output_path)
        else:
            raise ValueError(f"Unsupported content type: {content_type}")
    
    def _create_translated_pdf(self, original_content: Dict[str, Any], 
                             translated_texts: Dict[int, List[str]], 
                             output_path: str) -> str:
        """Create translated PDF with preserved formatting."""
        layouts = original_content['layouts']
        
        # Use the advanced PDF processor to create the translated document
        self.pdf_processor.create_translated_pdf(layouts, translated_texts, output_path)
        
        logger.info(f"Translated PDF created: {output_path}")
        return output_path
    
    def _create_translated_docx(self, original_content: Dict[str, Any], 
                              translated_texts: Dict[int, List[str]], 
                              output_path: str) -> str:
        """Create translated DOCX document."""
        from docx import Document
        from docx.shared import Inches
        
        # Create new document
        doc = Document()
        
        # Add translated paragraphs
        page_texts = translated_texts.get(0, [])
        original_paragraphs = original_content['paragraphs']
        
        for i, para_info in enumerate(original_paragraphs):
            # Use translated text if available
            if i < len(page_texts) and page_texts[i].strip():
                text = page_texts[i]
            else:
                text = para_info['text']
            
            # Add paragraph with original style
            paragraph = doc.add_paragraph(text)
            
            # Try to preserve some formatting
            if para_info['style'] != 'Normal':
                try:
                    paragraph.style = para_info['style']
                except:
                    pass  # Style might not exist in new document
        
        # Save document
        doc.save(output_path)
        
        logger.info(f"Translated DOCX created: {output_path}")
        return output_path
    
    def _create_translated_txt(self, original_content: Dict[str, Any], 
                             translated_texts: Dict[int, List[str]], 
                             output_path: str) -> str:
        """Create translated TXT document."""
        page_texts = translated_texts.get(0, [])
        original_lines = original_content['lines']
        
        translated_lines = []
        for i, original_line in enumerate(original_lines):
            # Use translated text if available
            if i < len(page_texts) and page_texts[i].strip():
                translated_lines.append(page_texts[i])
            else:
                translated_lines.append(original_line)
        
        # Write translated content
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(translated_lines))
        
        logger.info(f"Translated TXT created: {output_path}")
        return output_path
    
    def convert_format(self, input_path: str, target_format: str) -> str:
        """Convert document to different format."""
        input_ext = Path(input_path).suffix.lower()
        target_ext = f'.{target_format.lower()}'
        
        if input_ext == target_ext:
            return input_path
        
        output_path = input_path.replace(input_ext, target_ext)
        
        # Simple format conversion (can be enhanced)
        if input_ext == '.pdf' and target_ext == '.txt':
            return self._pdf_to_txt(input_path, output_path)
        elif input_ext == '.docx' and target_ext == '.txt':
            return self._docx_to_txt(input_path, output_path)
        elif input_ext == '.txt' and target_ext == '.pdf':
            return self._txt_to_pdf(input_path, output_path)
        else:
            logger.warning(f"Format conversion not supported: {input_ext} -> {target_ext}")
            return input_path
    
    def _pdf_to_txt(self, pdf_path: str, txt_path: str) -> str:
        """Convert PDF to TXT."""
        doc = fitz.open(pdf_path)
        text_content = ""
        
        for page in doc:
            text_content += page.get_text() + "\n\n"
        
        doc.close()
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        return txt_path
    
    def _docx_to_txt(self, docx_path: str, txt_path: str) -> str:
        """Convert DOCX to TXT."""
        doc = Document(docx_path)
        text_content = ""
        
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        
        return txt_path
    
    def _txt_to_pdf(self, txt_path: str, pdf_path: str) -> str:
        """Convert TXT to PDF."""
        with open(txt_path, 'r', encoding='utf-8') as f:
            text_content = f.read()
        
        doc = fitz.open()
        # Split text into lines and handle pagination
        lines = text_content.split('\n')
        page = doc.new_page()
        y_position = 72
        line_height = 14
        page_height = page.rect.height - 72  # Bottom margin
        
        for line in lines:
            if y_position + line_height > page_height:
                page = doc.new_page()
                y_position = 72
            
            page.insert_text(
                (72, y_position),
                line,
                fontname="helv",
                fontsize=12,
                color=(0, 0, 0)
            )
            y_position += line_height
        
        doc.save(pdf_path)
        doc.close()
        
        return pdf_path
    def generate_preview(self, file_path: str, max_chars: int = 1000) -> str:
        """Generate a preview of the document content."""
        try:
            content = self.extract_content(file_path)
            
            if content['type'] == 'pdf_advanced':
                return content['preview']
            else:
                preview_text = content.get('preview', content.get('text_content', ''))
                if len(preview_text) > max_chars:
                    return preview_text[:max_chars] + '...'
                return preview_text
                
        except Exception as e:
            logger.error(f"Error generating preview: {e}")
            return f"Error generating preview: {str(e)}"
    
    def _get_backup_path(self, original_path: str) -> str:
        """Generate backup path for layout information."""
        backup_dir = os.path.join(os.path.dirname(original_path), '.layout_backups')
        os.makedirs(backup_dir, exist_ok=True)
        
        filename = Path(original_path).stem
        return os.path.join(backup_dir, f"{filename}_layout.json")