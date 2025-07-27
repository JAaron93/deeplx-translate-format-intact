"""Document processing service for multiple file formats"""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any

import fitz  # PyMuPDF
from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing for multiple formats"""

    def __init__(self) -> None:
        self.supported_formats: list[str] = [".pdf", ".docx", ".txt"]

    def extract_content(self, file_path: str) -> dict[str, Any]:
        """Extract content from document based on file type"""
        file_ext = Path(file_path).suffix.lower()

        if file_ext == ".pdf":
            return self._extract_pdf_content(file_path)
        elif file_ext == ".docx":
            return self._extract_docx_content(file_path)
        elif file_ext == ".txt":
            return self._extract_txt_content(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

    def _extract_pdf_content(self, file_path: str) -> dict[str, Any]:
        """Extract content from PDF file"""
        try:
            doc = fitz.open(file_path)
            pages = []

            for page_num in range(len(doc)):
                page = doc[page_num]

                # Extract text blocks with formatting
                blocks = page.get_text("dict")
                text_blocks = []

                for block in blocks["blocks"]:
                    if "lines" in block:
                        for line in block["lines"]:
                            for span in line["spans"]:
                                text_blocks.append(
                                    {
                                        "text": span["text"],
                                        "formatting": {
                                            "font": span["font"],
                                            "size": span["size"],
                                            "flags": span["flags"],
                                            "color": span["color"],
                                        },
                                        "position": {"bbox": span["bbox"]},
                                    }
                                )

                # Extract images
                images = []
                image_list = page.get_images(full=True)
                for img_index, img in enumerate(image_list):
                    pix = None  # Ensure defined for cleanup
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(doc, xref)
                        if pix.n - pix.alpha < 4:
                            image_rects = page.get_image_rects(xref)
                            images.append(
                                {
                                    "index": img_index,
                                    "xref": xref,
                                    "bbox": image_rects[0] if image_rects else None,
                                    "data": pix.tobytes("png"),
                                }
                            )
                    except Exception as e:
                        logger.warning(f"Could not extract image {img_index}: {e}")
                    finally:
                        # Explicitly dispose Pixmap to free memory
                        if pix is not None:
                            try:
                                pix.dispose()  # type: ignore[attr-defined]
                            except AttributeError:
                                pass
                            pix = None
                pages.append(
                    {
                        "page_num": page_num,
                        "text_blocks": text_blocks,
                        "images": images,
                        "size": (page.rect.width, page.rect.height),
                    }
                )

            doc.close()

            return {
                "type": "pdf",
                "pages": pages,
                "metadata": {"total_pages": len(pages)},
            }

        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

    def _extract_docx_content(self, file_path: str) -> dict[str, Any]:
        """Extract content from DOCX file"""
        try:
            doc = Document(file_path)

            # Extract paragraphs with formatting
            text_blocks = []
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    # Get paragraph formatting
                    formatting = {
                        "style": paragraph.style.name if paragraph.style else "Normal",
                        "alignment": paragraph.alignment,
                    }

                    # Get run-level formatting
                    runs = []
                    for run in paragraph.runs:
                        runs.append(
                            {
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline,
                                "font_name": run.font.name,
                                "font_size": run.font.size,
                            }
                        )

                    text_blocks.append(
                        {
                            "text": paragraph.text,
                            "formatting": formatting,
                            "runs": runs,
                            "position": {"paragraph": para_num},
                        }
                    )

            # Extract images
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    images.append(
                        {
                            "relationship_id": rel.rId,
                            "target": rel.target_ref,
                            "data": rel.target_part.blob,
                        }
                    )

            return {
                "type": "docx",
                "pages": [
                    {"page_num": 0, "text_blocks": text_blocks, "images": images}
                ],
                "metadata": {"total_pages": 1, "paragraphs": len(text_blocks)},
            }

        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise

    def _extract_txt_content(self, file_path: str) -> dict[str, Any]:
        """Extract content from TXT file"""
        try:
            # Check file size before reading
            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB limit for text files
            if file_size > max_size:
                raise ValueError(f"Text file too large: {file_size} bytes")

            with open(file_path, encoding="utf-8") as f:
                content = f.read()

            # Split into paragraphs
            paragraphs = content.split("\n\n")
            text_blocks = []

            for para_num, paragraph in enumerate(paragraphs):
                if paragraph.strip():
                    text_blocks.append(
                        {
                            "text": paragraph.strip(),
                            "formatting": {"style": "plain"},
                            "position": {"paragraph": para_num},
                        }
                    )

            return {
                "type": "txt",
                "pages": [{"page_num": 0, "text_blocks": text_blocks, "images": []}],
                "metadata": {"total_pages": 1, "paragraphs": len(text_blocks)},
            }

        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            raise

    def generate_preview(self, file_path: str, max_chars: int = 1000) -> str:
        """Generate a preview of the document content"""
        try:
            content = self.extract_content(file_path)

            preview_text = ""
            char_count = 0

            for page in content["pages"]:
                for block in page["text_blocks"]:
                    text = block["text"]
                    if char_count + len(text) > max_chars:
                        preview_text += text[: max_chars - char_count] + "..."
                        break
                    preview_text += text + "\n"
                    char_count += len(text) + 1

                if char_count >= max_chars:
                    break

            return preview_text

        except Exception as e:
            logger.error(f"Error generating preview for {file_path}: {e}")
            return ""

    def create_translated_document(
        self,
        translated_content: dict[str, Any],
        output_filename: str,
        original_format: str,
    ) -> str:
        """Create translated document in specified format"""
        output_dir = "downloads"
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, output_filename)

        if original_format == ".pdf":
            return self._create_pdf_document(translated_content, output_path)
        elif original_format == ".docx":
            return self._create_docx_document(translated_content, output_path)
        elif original_format == ".txt":
            return self._create_txt_document(translated_content, output_path)
        else:
            raise ValueError(f"Unsupported output format: {original_format}")

    def _create_pdf_document(self, content: dict[str, Any], output_path: str) -> str:
        """Create PDF document with translated content"""
        try:
            doc = fitz.open()

            for page_data in content["pages"]:
                # Create new page
                page = doc.new_page(
                    width=page_data.get("size", (595, 842))[0],
                    height=page_data.get("size", (595, 842))[1],
                )

                # Add text blocks
                for block in page_data["text_blocks"]:
                    try:
                        # Get position and formatting
                        bbox = block.get("position", {}).get("bbox", [50, 50, 500, 100])
                        formatting = block.get("formatting", {})

                        # Insert text
                        page.insert_textbox(
                            fitz.Rect(bbox),
                            block["text"],
                            fontname=formatting.get("font", "helv"),
                            fontsize=formatting.get("size", 12),
                            color=self._normalize_color(formatting.get("color", 0)),
                        )
                    except Exception as e:
                        logger.warning(f"Could not insert text block: {e}")

                # Add images
                for image in page_data.get("images", []):
                    try:
                        if image.get("bbox") and image.get("data"):
                            page.insert_image(
                                fitz.Rect(image["bbox"]), stream=image["data"]
                            )
                    except Exception as e:
                        logger.warning(f"Could not insert image: {e}")

            doc.save(output_path)
            doc.close()

            return output_path

        except Exception as e:
            logger.error(f"PDF creation error: {e}")
            raise

    def _create_docx_document(self, content: dict[str, Any], output_path: str) -> str:
        """Create DOCX document with translated content"""
        try:
            doc = Document()

            for page_data in content["pages"]:
                for block in page_data["text_blocks"]:
                    paragraph = doc.add_paragraph()

                    # Apply formatting if available
                    formatting = block.get("formatting", {})
                    if "style" in formatting:
                        try:
                            paragraph.style = formatting["style"]
                        except:
                            pass

                    # Add text with run-level formatting
                    if "runs" in block:
                        for run_data in block["runs"]:
                            run = paragraph.add_run(run_data["text"])
                            if run_data.get("bold"):
                                run.bold = True
                            if run_data.get("italic"):
                                run.italic = True
                            if run_data.get("underline"):
                                run.underline = True
                    else:
                        paragraph.add_run(block["text"])

            doc.save(output_path)
            return output_path

        except Exception as e:
            logger.error(f"DOCX creation error: {e}")
            raise

    def _create_txt_document(self, content: dict[str, Any], output_path: str) -> str:
        """Create TXT document with translated content"""
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                for page_data in content["pages"]:
                    for block in page_data["text_blocks"]:
                        f.write(block["text"] + "\n\n")

            return output_path

        except Exception as e:
            logger.error(f"TXT creation error: {e}")
            raise

    def _normalize_color(self, color_value: Any) -> tuple[float, float, float]:
        """Convert color to RGB tuple"""
        if isinstance(color_value, int):
            r = (color_value >> 16) & 0xFF
            g = (color_value >> 8) & 0xFF
            b = color_value & 0xFF
            return (r / 255.0, g / 255.0, b / 255.0)
        return (0, 0, 0)

    def convert_format(self, input_file: str, target_format: str) -> str:
        """Convert document to different format"""
        try:
            # Extract content from input file
            content = self.extract_content(input_file)

            # Generate new filename
            input_path = Path(input_file)
            output_filename = f"{input_path.stem}.{target_format}"

            # Create document in target format
            return self.create_translated_document(
                content, output_filename, f".{target_format}"
            )

        except Exception as e:
            logger.error(f"Format conversion error: {e}")
            raise
